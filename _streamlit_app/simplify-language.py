import re
import time
import base64
import io
import logging
import numpy as np
import os
import requests
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches
from zix.understandability import get_zix, get_cefr
from metrics import track_metrics
from utils_sample_texts import SAMPLE_TEXT_01
from utils_prompts import (
    SYSTEM_MESSAGE_EASIER,
    SYSTEM_MESSAGE_ES,
    SYSTEM_MESSAGE_LS,
    SYSTEM_MESSAGE_ANALYSIS,
    RULES_EASIER,
    RULES_ES,
    RULES_LS,
    REWRITE_COMPLETE,
    OPENAI_TEMPLATE_EASIER,
    OPENAI_TEMPLATE_ES,
    OPENAI_TEMPLATE_LS,
    OPENAI_TEMPLATE_ANALYSIS_GENERIC,
)

import streamlit as st

st.set_page_config(layout="wide")

OPENAI_TEMPLATES = [
    OPENAI_TEMPLATE_EASIER,
    OPENAI_TEMPLATE_ES,
    OPENAI_TEMPLATE_LS,
    OPENAI_TEMPLATE_ANALYSIS_GENERIC,
]

# Constants
MODEL_OPTIONS = {
    "Gemma 3": "https://sp000201-t5.kt.ktzh.ch",
    "Phi-4": "https://sp000201-t6.kt.ktzh.ch",
}

# # Model-specific temperature settings
# MODEL_TEMPERATURES = {
#     # "Gemma 3": 1.0,
#     "default": 0.5,  # Used for all other models
# }

TEMPERATURE = 0.2

# MODEL_NAME is now assigned dynamically based on user selection (see below)
MAX_TOKENS = 4096

# Height of the text areas for input and output.
TEXT_AREA_HEIGHT = 600

# Maximum number of characters for the input text.
# This is way below the context window sizes of the models.
# We can increase this. However, we found that users can work and validate better when we nudge to work with shorter texts.
MAX_CHARS_INPUT = 5_000


USER_WARNING = """Mit der KlartextZH-App kannst du Texte sprachlich vereinfachen. Dazu schicken wir deinen Text an einen AFI KI-Server, den wir im Kanton betreiben. Du kannst daher auch vertrauliche Daten eingeben. Bitte beachte: KI-Sprachmodelle machen Fehler. Die App liefert lediglich einen Entwurf. Überprüfe das Ergebnis immer und passe es an, wenn nötig. Gib uns jederzeit [Feedback](mailto:patrick.arnecke@statistik.ji.zh.ch). 🚀 Aktuelle App-Version ist v01.3.1 Die letzte Aktualisierung war am 24.3.2025."""


# Constants for the formatting of the Word document that can be downloaded.
FONT_WORDDOC = "Arial"
FONT_SIZE_HEADING = 12
FONT_SIZE_PARAGRAPH = 9
FONT_SIZE_FOOTER = 7
DEFAULT_OUTPUT_FILENAME = "Ergebnis.docx"
ANALYSIS_FILENAME = "Analyse.docx"


# Limits for the understandability score to determine if the text is easy, medium or hard to understand.
LIMIT_HARD = -2
LIMIT_MEDIUM = 0

DATETIME_FORMAT = "%Y-%m-%d %H:%M:%S"


# Functions




@st.cache_resource
def get_project_info():
    """Get markdown for project information that is shown in the expander section at the top of the app."""
    # Use the current file's directory as the base path
    try:
        import os

        with open(os.path.join(os.path.dirname(__file__), "utils_expander.md")) as f:
            return f.read()
    except FileNotFoundError:
        # Fallback message if file is not found
        return """
        # KlartextZH - Sprache einfach vereinfachen

        Mit dieser Applikation kannst du komplexe Texte in verständlichere Sprache umwandeln.
        
        Die App bietet drei Stufen der Vereinfachung:
        - Verständlichere Sprache (B2-Niveau)
        - Einfache Sprache (B1-A2-Niveau)
        - Leichte Sprache (A2-A1-Niveau)
        
        ### Image ###
        
        Weitere Informationen folgen.
        """


@st.cache_resource
def create_project_info(project_info):
    """Create expander for project info. Add the image in the middle of the content."""
    import os

    with st.expander("Detaillierte Informationen zum Projekt"):
        project_info = project_info.split("### Image ###")
        st.markdown(project_info[0], unsafe_allow_html=True)
        try:
            image_path = os.path.join(
                os.path.dirname(__file__), "zix_scores_validation_de.jpg"
            )
            st.image(image_path, use_container_width=True)
        except FileNotFoundError:
            st.info(
                "Visualisierung zur Textverständlichkeit konnte nicht geladen werden."
            )
        st.markdown(project_info[1], unsafe_allow_html=True)


def create_prompt(
    text,
    prompt_easy,
    prompt_es,
    prompt_ls,
    analysis_generic,
    analysis,
):
    """Create prompt and system message according the app settings."""
    completeness = REWRITE_COMPLETE

    if simplification_level == "Verständlichere Sprache":
        final_prompt = prompt_easy.format(
            completeness=completeness, rules=RULES_EASIER, prompt=text
        )
        system = SYSTEM_MESSAGE_EASIER
        if analysis:
            final_prompt = analysis_generic.format(prompt=text)
            system = SYSTEM_MESSAGE_ANALYSIS
    elif simplification_level == "Einfache Sprache":
        final_prompt = prompt_es.format(
            completeness=completeness, rules=RULES_ES, prompt=text
        )
        system = SYSTEM_MESSAGE_ES
        if analysis:
            final_prompt = analysis_generic.format(prompt=text)
        system = SYSTEM_MESSAGE_ANALYSIS
    else:
        final_prompt = prompt_ls.format(
            completeness=completeness, rules=RULES_LS, prompt=text
        )
        system = SYSTEM_MESSAGE_LS
        if analysis:
            final_prompt = analysis_generic.format(prompt=text)
        system = SYSTEM_MESSAGE_ANALYSIS
    return final_prompt, system


def call_llm(
    text,
    model_name,
    analysis=False,
):
    """Call llama.cpp API using requests for text generation."""
    text = text.strip()
    final_prompt, system = create_prompt(text, *OPENAI_TEMPLATES, analysis)
    
    # Combine system and user prompt for llama.cpp
    full_prompt = f"{system}\n\n{final_prompt}"

    try:
        headers = {
            "Content-Type": "application/json",
        }
        
        payload = {
            "prompt": full_prompt,
            "n_predict": MAX_TOKENS,
            "temperature": TEMPERATURE,
        }

        timeout_value = 180 if analysis else 60

        # The URL is the model_name
        response = requests.post(
            f"{model_name}/completion",
            headers=headers,
            json=payload,
            timeout=timeout_value,
            verify=False  # Accept self-signed certificates
        )
        response.raise_for_status()

        # Extract the response message
        message = response.json().get("content", "")
        message = strip_markdown(message)

        return True, message

    except requests.exceptions.RequestException as e:
        print(f"Error: {e}")
        if isinstance(e, requests.exceptions.Timeout):
            return (
                False,
                f"Zeitüberschreitung: Die Anfrage dauerte länger als {timeout_value} Sekunden.",
            )
        return False, str(e)
    except Exception as e:
        print(f"Error: {e}")
        return False, str(e)


def get_result_from_response(response):
    """Extract text between tags from response."""
    if simplification_level == "Verständlichere Sprache":
        tag = "verständlichesprache"
    elif simplification_level == "Einfache Sprache":
        tag = "einfachesprache"
    else:
        tag = "leichtesprache"
    result = re.findall(rf"<{tag}>(.*?)</{tag}>", response, re.DOTALL)

    # Handle case when no matching tags are found
    if not result:
        print(
            f"Warning: Response was not properly formatted between <{tag}> tags. Just returning the whole response."
        )
        return response.strip()

    return "\n".join(result).strip()


def strip_markdown(text):
    """Strip markdown from text."""
    # Remove markdown headers
    text = re.sub(r"#+\s", "", text)
    # Remove markdown italic and bold
    text = re.sub(r"\*\*|\*|__|_", "", text)
    return text


def enter_sample_text():
    """Enter sample text into the text input in the left column."""
    st.session_state.key_textinput = SAMPLE_TEXT_01


def create_download_link(text_input, response, analysis=False):
    """Create a downloadable Word document and download link of the results."""
    document = Document()

    h1 = document.add_heading("Ausgangstext")
    p1 = document.add_paragraph("\n" + text_input)

    if analysis:
        h2 = document.add_heading(f"Analyse ({model_choice})")
    else:
        h2 = document.add_heading(f"Vereinfachter Text ({model_choice})")

    p2 = document.add_paragraph(response)
    timestamp = datetime.now().strftime(DATETIME_FORMAT)
    models_used = model_choice
    footer = document.sections[0].footer
    footer.paragraphs[
        0
    ].text = f"Erstellt am {timestamp} mit der KlartextZH-App » des Kantons Zürich.\nModell: {models_used}\nVerarbeitungszeit: {time_processed:.1f} Sekunden"

    # Set font for all paragraphs.
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT_WORDDOC

    # Set font size for all headings.
    for paragraph in [h1, h2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_HEADING)

    # Set font size for all paragraphs.
    for paragraph in [p1, p2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_PARAGRAPH)

    # Set font and font size for footer.
    for run in footer.paragraphs[0].runs:
        run.font.name = "Arial"
        run.font.size = Pt(FONT_SIZE_FOOTER)

    section = document.sections[0]
    section.page_width = Inches(8.27)  # Width of A4 paper in inches
    section.page_height = Inches(11.69)  # Height of A4 paper in inches

    io_stream = io.BytesIO()
    document.save(io_stream)

    # # A download button resets the app. So we use a link instead.
    # https://github.com/streamlit/streamlit/issues/4382#issuecomment-1223924851
    # https://discuss.streamlit.io/t/creating-a-pdf-file-generator/7613?u=volodymyr_holomb

    b64 = base64.b64encode(io_stream.getvalue())
    file_name = DEFAULT_OUTPUT_FILENAME
    caption = "Vereinfachten Text herunterladen"

    if analysis:
        caption = "Analyse herunterladen"
        file_name = ANALYSIS_FILENAME
    download_url = f'<a href="data:application/octet-stream;base64,{b64.decode()}" download="{file_name}">{caption}</a>'
    st.markdown(download_url, unsafe_allow_html=True)


def log_event(
    text,
    response,
    do_analysis,
    do_simplification,
    simplification_level,
    model_choice,
    time_processed,
    success,
):
    """Log event."""
    log_string = f"{datetime.now().strftime(DATETIME_FORMAT)}"
    log_string += f"\t{len(text.split())}"  # Number of words in the input text.
    log_string += f"\t{len(response.split())}"  # Number of words in the output text.
    log_string += f"\t{do_analysis}"
    log_string += f"\t{do_simplification}"
    log_string += f"\t{simplification_level}"
    log_string += f"\t{model_choice}"
    log_string += f"\t{time_processed:.3f}"
    log_string += f"\t{success}"

    logging.warning(log_string)

    # Also track metrics for Prometheus
    track_metrics(
        text,
        response,
        do_analysis,
        do_simplification,
        simplification_level,
        model_choice,
        time_processed,
        success,
    )


project_info = get_project_info()

if "key_textinput" not in st.session_state:
    st.session_state.key_textinput = ""

st.markdown("## 🙋‍♀️ KlartextZH - Sprache einfach vereinfachen")
create_project_info(project_info)
st.caption(USER_WARNING, unsafe_allow_html=True)
st.markdown("---")

# Set up first row with all buttons and settings.
button_cols = st.columns([1, 1, 1, 2])
with button_cols[0]:
    st.button(
        "Beispiel einfügen",
        on_click=enter_sample_text,
        use_container_width=True,
        help="Fügt einen Beispieltext ein.",
    )
with button_cols[1]:
    do_simplification = st.button(
        "Vereinfachen",
        use_container_width=True,
        help="Vereinfacht deinen Ausgangstext.",
    )
    do_analysis = st.button(
        "Analysieren",
        use_container_width=True,
        help="Analysiert deinen Ausgangstext Satz für Satz.",
    )


with button_cols[2]:
    simplification_level = st.radio(
        label="Grad der Vereinfachung",
        options=["Verständlichere Sprache", "Einfache Sprache", "Leichte Sprache"],
        index=1,
        help="**«Verständlichere Sprache»** überarbeitet den Text vorsichtiger und zielt auf Sprachniveau B2. **«Einfache Sprache»** formuliert nach den Regeln für Einfache Sprache und zielt auf B1 bis A2. **«Leichte Sprache»** folgt den Regeln für Leichte Sprache und zielt auf A2 bis A1.",
    )


with button_cols[3]:
    # Model selection
    # Model selection
    model_keys = list(MODEL_OPTIONS.keys())
    selected_model_key = st.radio(
        "Modell:",
        options=model_keys,
        index=0, # Default to the first one
        help="Wähle das Modell, das für die Vereinfachung verwendet werden soll."
    )
    # Get the actual URL for API calls
    MODEL_NAME = MODEL_OPTIONS[selected_model_key]
    # Store friendly name for display/logging
    model_choice = selected_model_key

cols = st.columns([2, 2, 1])

with cols[0]:
    source_text = st.container()
with cols[1]:
    placeholder_result = st.empty()
with cols[2]:
    placeholder_analysis = st.empty()

with source_text:
    st.text_area(
        "Ausgangstext, den du vereinfachen möchtest",
        value=None,
        height=TEXT_AREA_HEIGHT,
        max_chars=MAX_CHARS_INPUT,
        key="key_textinput",
    )
with placeholder_result:
    text_output = st.text_area(
        "Ergebnis",
        value=" ",
        height=TEXT_AREA_HEIGHT,
    )
with placeholder_analysis:
    text_analysis = st.metric(
        label="Verständlichkeit",
        value=None,
        delta=None,
        help="Texte in Einfacher Sprache haben meist einen Wert 0 oder höher, Texte in Leichter Sprache haben Werte von 2 oder höher.",
    )


# Client is now fetched dynamically within call_llm based on selected model

# Start processing if one of the processing buttons is clicked.
if do_simplification or do_analysis:
    start_time = time.time()
    if st.session_state.key_textinput == "":
        st.error("Bitte gib einen Text ein.")
        st.stop()

    score_source = get_zix(st.session_state.key_textinput)
    # We add 0 to avoid negative zero.
    score_source_rounded = int(np.round(score_source, 0) + 0)
    cefr_source = get_cefr(score_source)

    # Analyze source text and display results.
    with source_text:
        if score_source < LIMIT_HARD:
            st.markdown(
                f"Dein Ausgangstext ist **:red[schwer verständlich]**. (Wert: {score_source_rounded}). Das entspricht grob geschätzt dem **:red[Sprachniveau {cefr_source}]**."
            )
        elif score_source >= LIMIT_HARD and score_source < LIMIT_MEDIUM:
            st.markdown(
                f"Dein Ausgangstext ist **:orange[durchschnittlich verständlich]**. (Wert: {score_source_rounded}). Das entspricht grob geschätzt dem **:orange[Sprachniveau {cefr_source}]**."
            )
        else:
            st.markdown(
                f"Dein Ausgangstext ist **:green[gut verständlich]**. (Wert: {score_source_rounded}). Das entspricht grob geschätzt dem **:green[Sprachniveau {cefr_source}]**."
            )
        with placeholder_analysis.container():
            text_analysis = st.metric(
                label="Verständlichkeit",
                value=score_source_rounded,
                delta=None,
                help="Verständlichkeit auf einer Skala von -10 bis 10 Punkten (von -10 = extrem schwer verständlich bis 10 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 0 bis 4 oder höher, Texte in Leichter Sprache 2 bis 6 oder höher.",
            )

        with placeholder_analysis.container():
            with st.spinner("Text wird verarbeitet..."):
                success, response = call_llm(
                    st.session_state.key_textinput,
                    model_name=MODEL_NAME,
                    analysis=do_analysis,
                )

    if success is False:
        error_message = str(response).replace("'", "")[
            :200
        ]  # Limit length and remove quotes for display
        st.error(
            f"Es ist ein Fehler bei der Abfrage aufgetreten: {error_message}. Bitte versuche es erneut."
        )
        time_processed = time.time() - start_time
        log_event(
            st.session_state.key_textinput,
            "Error from model call.",
            do_analysis,
            do_simplification,
            simplification_level,
            model_choice,
            time_processed,
            success,
        )

        st.stop()

    # Display results in UI.
    text = "Dein vereinfachter Text"
    if do_analysis:
        text = "Deine Analyse"
    # Often the models return the German letter ß. Replace it with the Swiss German equivalent ss.
    response = response.replace("ß", "ss")
    time_processed = time.time() - start_time

    with placeholder_result.container():
        st.text_area(
            text,
            height=TEXT_AREA_HEIGHT,
            value=response,
        )
        if do_simplification:
            # Check if we have content to analyze
            if response.strip():
                # Get ZIX score and handle None values
                score_target = get_zix(response)
                # Only proceed with calculations if we got a valid score
                if score_target is not None:
                    score_target_rounded = int(np.round(score_target, 0) + 0)
                    cefr_target = get_cefr(score_target)
                else:
                    # Handle case where ZIX scoring returns None
                    score_target_rounded = 0
                    cefr_target = "?"
            else:
                # Handle empty response case
                score_target_rounded = 0
                cefr_target = "?"

            if score_target is not None and score_target < LIMIT_HARD:
                st.markdown(
                    f"Dein vereinfachter Text ist **:red[schwer verständlich]**. (Wert: {score_target_rounded}). Das entspricht etwa dem **:red[Sprachniveau {cefr_target}]**."
                )
            elif score_target >= LIMIT_HARD and score_target < LIMIT_MEDIUM:
                st.markdown(
                    f"Dein Ausgangstext ist **:orange[durchschnittlich verständlich]**. (Wert: {score_target_rounded}). Das entspricht grob geschätzt dem **:orange[Sprachniveau {cefr_target}]**."
                )
            else:
                st.markdown(
                    f"Dein vereinfachter Text ist **:green[gut verständlich]**. (Wert: {score_target_rounded}). Das entspricht etwa dem **:green[Sprachniveau {cefr_target}]**."
                )
            with placeholder_analysis.container():
                # Calculate delta safely only if score_target is valid
                if score_target is not None:
                    delta = int(np.round(score_target - score_source, 0))
                else:
                    delta = None

                text_analysis = st.metric(
                    label="Verständlichkeit",
                    value=score_target_rounded,
                    delta=delta,
                    help="Verständlichkeit auf einer Skala von -10 bis 10 (von -10 = extrem schwer verständlich bis 10 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 0 bis 4 oder höher.",
                )

                create_download_link(st.session_state.key_textinput, response)
                st.caption(f"Verarbeitet in {time_processed:.1f} Sekunden.")
        else:
            with placeholder_analysis.container():
                text_analysis = st.metric(
                    label="Verständlichkeit 0-20",
                    value=score_source_rounded,
                    help="Verständlichkeit auf einer Skala von -10 bis 10 (von -10 = extrem schwer verständlich bis 10 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 0 bis 4 oder höher.",
                )
                create_download_link(
                    st.session_state.key_textinput, response, analysis=True
                )
                st.caption(f"Verarbeitet in {time_processed:.1f} Sekunden.")

        log_event(
            st.session_state.key_textinput,
            response,
            do_analysis,
            do_simplification,
            simplification_level,
            model_choice,
            time_processed,
            success,
        )
        st.stop()
