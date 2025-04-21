# ---------------------------------------------------------------
# Imports

import streamlit as st

st.set_page_config(layout="wide")

import time
import base64
import io
import logging
import numpy as np
import os
from datetime import datetime
from openai import OpenAI

from docx import Document
from docx.shared import Pt, Inches
from zix.understandability import get_zix, get_cefr

# from metrics import track_metrics
from utils_prompts import (
    SYSTEM_MESSAGE_VS,
    SYSTEM_MESSAGE_ES,
    SYSTEM_MESSAGE_LS,
    ADDITION_GEMMA,
)


# ---------------------------------------------------------------
# Constants

DEFAULT_SYSTEM_MESSAGE = SYSTEM_MESSAGE_VS

MODEL_OPTIONS = {
    "Phi-4 KTZH": "hf.co/matterhorn/Phi-4_KTZH_250422_v1:latest",
    "Gemma 3": "gemma3:27b-it-qat",
}

# Model-specific temperature settings
MODEL_TEMPERATURES = {
    "Gemma 3": 0.5,
    "Phi-4 KTZH": 0.5,
}

DEFAULT_MODEL = "Phi-4 KTZH"
DEFAULT_TEMPERATURE = 0.5


# Get Ollama host from environment variable or use default
OLLAMA_HOST = os.environ.get("OLLAMA_HOST", "http://localhost:11434")
MAX_TOKENS = 2048

# Height of the text areas for input and output.
TEXT_AREA_HEIGHT = 600

# Maximum number of characters for the input text.
# This is way below the context window sizes of the models.
# We can increase this. However, we found that users can work and validate better when we nudge to work with shorter texts.
MAX_CHARS_INPUT = 10_000


USER_WARNING = """Mit der KlartextZH-App kannst du Texte sprachlich vereinfachen. Dazu schicken wir deinen Text an den AFI KI-Server, den wir im Kanton betreiben. Du kannst daher auch vertrauliche Daten eingeben. Bitte beachte: KI-Sprachmodelle machen Fehler. Die App liefert lediglich einen Entwurf. Überprüfe das Ergebnis immer und passe es an, wenn nötig. Gib uns jederzeit [Feedback](mailto:patrick.arnecke@statistik.ji.zh.ch). 🚀 Aktuelle App-Version ist v.02. Die letzte Aktualisierung war am 22.4.2025."""


# Constants for the formatting of the Word document that can be downloaded.
FONT_WORDDOC = "Arial"
FONT_SIZE_HEADING = 12
FONT_SIZE_PARAGRAPH = 9
FONT_SIZE_FOOTER = 7
DEFAULT_OUTPUT_FILENAME = "Ergebnis.docx"
ANALYSIS_FILENAME = "Analyse.docx"


# Limits for the understandability score to determine if the text is easy, medium or hard to understand.
LIMIT_HARD = -2
LIMIT_MEDIUM = 0

DATETIME_FORMAT = "%Y-%m-%d %H:%M:%S"


# ---------------------------------------------------------------
# Functions


@st.cache_resource
def get_ollama_client(base_url=OLLAMA_HOST):
    """Create a connection to the Ollama API using OpenAI SDK.

    Args:
        base_url (str): URL of the Ollama service.

    Returns:
        OpenAI: Configured OpenAI client for Ollama.
    """
    try:
        client = OpenAI(
            base_url=f"{base_url}/v1",
            api_key="ollama",  # Required but unused for Ollama
        )
        return client
    except Exception as e:
        st.error(f"Verbindung zum AFI KI Server fehlgeschlagen: {str(e)}")
        return None


@st.cache_resource
def get_project_info():
    """Get markdown for project information that is shown in the expander section."""
    with open("utils_expander.md") as f:
        return f.read()


@st.cache_resource
def create_project_info(project_info):
    """Create expander for project info. Add the image in the middle of the content."""
    with st.expander("Detaillierte Informationen zum Projekt"):
        project_info = project_info.split("### Image ###")
        st.markdown(project_info[0], unsafe_allow_html=True)
        try:
            image_path = os.path.join(
                os.path.dirname(__file__), "zix_scores_validation_de.jpg"
            )
            st.image(image_path, use_container_width=True)
        except FileNotFoundError:
            st.info(
                "Visualisierung zur Textverständlichkeit konnte nicht geladen werden."
            )
        st.markdown(project_info[1], unsafe_allow_html=True)


def call_llm(
    client,
    text,
    model_name=DEFAULT_MODEL,
    system_message=DEFAULT_SYSTEM_MESSAGE,
):
    """Call Ollama API using OpenAI SDK for text generation."""
    if not client:
        return False, "Verbindung zum KI Server nicht verfügbar."

    temperature = MODEL_TEMPERATURES.get(model_name, DEFAULT_TEMPERATURE)
    model_id = MODEL_OPTIONS.get(model_name)

    try:
        response = client.chat.completions.create(
            model=model_id,
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": text.strip()},
            ],
            temperature=temperature,
            max_tokens=MAX_TOKENS,
        )
        return True, response.choices[0].message.content.strip()

    except Exception as e:
        if "timeout" in str(e).lower():
            return (
                False,
                f"Zeitüberschreitung: Die Anfrage dauerte zu lange und hat nicht funktioniert. Bitte versuche es erneut.",
            )
        return False, str(e)


def create_download_link(text_input, response, selected_model, time_processed):
    """Create a downloadable Word document and download link of the results."""
    document = Document()

    h1 = document.add_heading("Ausgangstext")
    p1 = document.add_paragraph("\n" + text_input)
    h2 = document.add_heading(f"Vereinfachter Text ({selected_model})")
    p2 = document.add_paragraph(response)
    timestamp = datetime.now().strftime(DATETIME_FORMAT)
    footer = document.sections[0].footer
    footer.paragraphs[
        0
    ].text = f"Erstellt am {timestamp} mit der KlartextZH-App » des Kantons Zürich.\nModell: {selected_model}\nVerarbeitungszeit: {time_processed:.1f} Sekunden"

    # Set font for all paragraphs.
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT_WORDDOC

    # Set font size for all headings.
    for paragraph in [h1, h2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_HEADING)

    # Set font size for all paragraphs.
    for paragraph in [p1, p2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_PARAGRAPH)

    # Set font and font size for footer.
    for run in footer.paragraphs[0].runs:
        run.font.name = "Arial"
        run.font.size = Pt(FONT_SIZE_FOOTER)

    section = document.sections[0]
    section.page_width = Inches(8.27)  # Width of A4 paper in inches
    section.page_height = Inches(11.69)  # Height of A4 paper in inches

    io_stream = io.BytesIO()
    document.save(io_stream)

    # # A download button resets the app. So we use a link instead.
    # https://github.com/streamlit/streamlit/issues/4382#issuecomment-1223924851
    # https://discuss.streamlit.io/t/creating-a-pdf-file-generator/7613?u=volodymyr_holomb

    b64 = base64.b64encode(io_stream.getvalue())
    file_name = DEFAULT_OUTPUT_FILENAME
    caption = "Vereinfachten Text herunterladen"

    download_url = f'<a href="data:application/octet-stream;base64,{b64.decode()}" download="{file_name}">{caption}</a>'
    st.markdown(download_url, unsafe_allow_html=True)


def log_event(
    text,
    response,
    do_simplification,
    simplification_level,
    selected_model,
    time_processed,
    success,
):
    """Log event."""
    log_string = f"{datetime.now().strftime(DATETIME_FORMAT)}"
    log_string += f"\t{len(text.split())}"  # Number of words in the input text.
    log_string += f"\t{len(response.split())}"  # Number of words in the output text.
    log_string += f"\t{do_simplification}"
    log_string += f"\t{simplification_level}"
    log_string += f"\t{selected_model}"
    log_string += f"\t{time_processed:.3f}"
    log_string += f"\t{success}"

    logging.info(log_string)

    # # Also track metrics for Prometheus
    # track_metrics(
    #     text,
    #     response,
    #     do_simplification,
    #     simplification_level,
    #     selected_model,
    #     time_processed,
    #     success,
    # )


# ---------------------------------------------------------------
# Main App

ollama_client = get_ollama_client()
project_info = get_project_info()

cols = st.columns([1, 7])
st.markdown("## 🙋‍♀️ KlartextZH - Sprache einfach vereinfachen", unsafe_allow_html=True)
create_project_info(project_info)
st.caption(USER_WARNING, unsafe_allow_html=True)
st.markdown("---")

# Set up first row with buttons and settings.
button_cols = st.columns([1, 3, 1])

with button_cols[0]:
    # Disable button if client connection failed
    simplify_disabled = ollama_client is None
    do_simplification = st.button(
        "Vereinfachen",
        use_container_width=True,
        help="Vereinfacht deinen Ausgangstext.",
        disabled=simplify_disabled,
    )

with button_cols[1]:
    simplification_level = st.radio(
        label="Grad der Vereinfachung",
        options=["Verständlichere Sprache", "Einfache Sprache", "Leichte Sprache"],
        index=0,
        help="**«Verständlichere Sprache»** überarbeitet den Text vorsichtiger und zielt auf Sprachniveau B2. **«Einfache Sprache»** folgt den Regeln für Einfache Sprache (B1 bis A2) und **«Leichte Sprache»** folgt den Regeln für Leichte Sprache (A2 bis A1).",
        horizontal=True,
    )

with button_cols[2]:
    selected_model = st.radio(
        "Modell:",
        options=list(MODEL_OPTIONS.keys()),
        index=0,
        help="Wähle das Modell, das für die Vereinfachung verwendet werden soll. Phi-4 KTZH ist das Standardmodell und von uns für die Anwendung im Kanton optimiert. Gemma-3 ist ein gutes allgemeines Modell, das auch einen Versuch wert ist.",
    )


system_messages = {
    "Verständlichere Sprache": SYSTEM_MESSAGE_VS,
    "Einfache Sprache": SYSTEM_MESSAGE_ES,
    "Leichte Sprache": SYSTEM_MESSAGE_LS,
}

system_message = system_messages[simplification_level]
if selected_model == "Gemma 3":
    system_message += ADDITION_GEMMA

cols = st.columns([2, 2, 1])

with cols[0]:
    source_text = st.container()
with cols[1]:
    placeholder_result = st.empty()
with cols[2]:
    placeholder_analysis = st.empty()

with source_text:
    st.text_area(
        "Ausgangstext, den du vereinfachen möchtest",
        value=None,
        height=TEXT_AREA_HEIGHT,
        max_chars=MAX_CHARS_INPUT,
        key="key_textinput",
    )
with placeholder_result:
    text_output = st.text_area(
        "Ergebnis",
        value=" ",
        height=TEXT_AREA_HEIGHT,
    )
with placeholder_analysis:
    text_analysis = st.metric(
        label="Verständlichkeit",
        value=None,
        delta=None,
        help="Texte in Einfacher Sprache haben meist einen Wert 0 oder höher, Texte in Leichter Sprache haben Werte von 2 oder höher.",
    )


# Start processing if one of the processing buttons is clicked.
if do_simplification:
    start_time = time.time()
    if st.session_state.key_textinput == "":
        st.error("Bitte gib einen Text ein.")
        st.stop()

    score_source = get_zix(st.session_state.key_textinput)
    # We add 0 to avoid negative zero.
    score_source_rounded = int(np.round(score_source, 0) + 0)
    cefr_source = get_cefr(score_source)

    # Analyze source text and display results.
    with source_text:
        if score_source < LIMIT_HARD:
            st.markdown(
                f"Dein Ausgangstext ist **:red[schwer verständlich]**. (Wert: {score_source_rounded}). Das entspricht grob geschätzt dem **:red[Sprachniveau {cefr_source}]**."
            )
        elif score_source >= LIMIT_HARD and score_source < LIMIT_MEDIUM:
            st.markdown(
                f"Dein Ausgangstext ist **:orange[durchschnittlich verständlich]**. (Wert: {score_source_rounded}). Das entspricht grob geschätzt dem **:orange[Sprachniveau {cefr_source}]**."
            )
        else:
            st.markdown(
                f"Dein Ausgangstext ist **:green[gut verständlich]**. (Wert: {score_source_rounded}). Das entspricht grob geschätzt dem **:green[Sprachniveau {cefr_source}]**."
            )
        with placeholder_analysis.container():
            text_analysis = st.metric(
                label="Verständlichkeit",
                value=score_source_rounded,
                delta=None,
                help="Verständlichkeit auf einer Skala von -10 bis 10 Punkten (von -10 = extrem schwer verständlich bis 10 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 0 bis 4 oder höher, Texte in Leichter Sprache 2 bis 6 oder höher.",
            )

        with placeholder_analysis.container():
            with st.spinner("Text wird verarbeitet..."):
                success, response = call_llm(
                    client=ollama_client,
                    text=st.session_state.key_textinput,
                    model_name=selected_model,
                    system_message=system_message,
                )

    if success is False:
        error_message = str(response).replace("'", "")[
            :200
        ]  # Limit length and remove quotes for display
        st.error(
            f"Es ist ein Fehler bei der Abfrage aufgetreten: {error_message}. Bitte versuche es erneut."
        )
        time_processed = time.time() - start_time
        log_event(
            st.session_state.key_textinput,
            "Error from model call.",
            do_simplification,
            simplification_level,
            selected_model,
            time_processed,
            success,
        )

        st.stop()

    # Display results in UI.
    text = "Dein vereinfachter Text"
    # Often the models return the German letter ß. Replace it with the Swiss German equivalent ss.
    response = response.replace("ß", "ss")
    time_processed = time.time() - start_time

    with placeholder_result.container():
        st.text_area(
            text,
            height=TEXT_AREA_HEIGHT,
            value=response,
        )
        if do_simplification:
            # Check if we have content to analyze
            if response.strip():
                # Get ZIX score and handle None values
                score_target = get_zix(response)
                # Only proceed with calculations if we got a valid score
                if score_target is not None:
                    score_target_rounded = int(np.round(score_target, 0) + 0)
                    cefr_target = get_cefr(score_target)
                else:
                    # Handle case where ZIX scoring returns None
                    score_target_rounded = 0
                    cefr_target = "?"
            else:
                # Handle empty response case
                score_target_rounded = 0
                cefr_target = "?"

            if score_target is not None and score_target < LIMIT_HARD:
                st.markdown(
                    f"Dein vereinfachter Text ist **:red[schwer verständlich]**. (Wert: {score_target_rounded}). Das entspricht etwa dem **:red[Sprachniveau {cefr_target}]**."
                )
            elif score_target >= LIMIT_HARD and score_target < LIMIT_MEDIUM:
                st.markdown(
                    f"Dein Ausgangstext ist **:orange[durchschnittlich verständlich]**. (Wert: {score_target_rounded}). Das entspricht grob geschätzt dem **:orange[Sprachniveau {cefr_target}]**."
                )
            else:
                st.markdown(
                    f"Dein vereinfachter Text ist **:green[gut verständlich]**. (Wert: {score_target_rounded}). Das entspricht etwa dem **:green[Sprachniveau {cefr_target}]**."
                )
            with placeholder_analysis.container():
                # Calculate delta safely only if score_target is valid
                if score_target is not None:
                    delta = int(np.round(score_target - score_source, 0))
                else:
                    delta = None

                text_analysis = st.metric(
                    label="Verständlichkeit",
                    value=score_target_rounded,
                    delta=delta,
                    help="Verständlichkeit auf einer Skala von -10 bis 10 (von -10 = extrem schwer verständlich bis 10 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 0 bis 4 oder höher.",
                )

                create_download_link(st.session_state.key_textinput, response, selected_model, time_processed)
                st.caption(f"Verarbeitet in {time_processed:.1f} Sekunden.")
        else:
            with placeholder_analysis.container():
                text_analysis = st.metric(
                    label="Verständlichkeit 0-20",
                    value=score_source_rounded,
                    help="Verständlichkeit auf einer Skala von -10 bis 10 (von -10 = extrem schwer verständlich bis 10 = sehr gut verständlich). Texte in Einfacher Sprache haben meist einen Wert von 0 bis 4 oder höher.",
                )
                create_download_link(
                    st.session_state.key_textinput, response
                )
                st.caption(f"Verarbeitet in {time_processed:.1f} Sekunden.")

        log_event(
            st.session_state.key_textinput,
            response,
            do_simplification,
            simplification_level,
            selected_model,
            time_processed,
            success,
        )
        st.stop()
