# ---------------------------------------------------------------
# Imports

import streamlit as st

st.set_page_config(
    layout="wide", page_title="KlartextZH - Sprache einfach vereinfachen"
)

import time
import base64
import io
import logging
import numpy as np
import os
from datetime import datetime
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.shared import Pt, Inches
from zix.understandability import get_zix, get_cefr
from metrics import track_metrics
import re
from utils_prompts import (
    SYSTEM_MESSAGE_VS,
    SYSTEM_MESSAGE_ES,
    SYSTEM_MESSAGE_LS,
    ADDITION_GEMMA,
)

logging.basicConfig(
    filename="app.log",
    level=logging.INFO,
    format="%(asctime)s\t%(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)


# ---------------------------------------------------------------
# Constants

# Use a .env file to manage environment variables
load_dotenv()

# Get Llama.cpp server URLs from environment variables
LLAMACPP_GPU0_URL = os.getenv("LLAMACPP_GPU0_URL")
LLAMACPP_GPU1_URL = os.getenv("LLAMACPP_GPU1_URL")

DEFAULT_SYSTEM_MESSAGE = SYSTEM_MESSAGE_VS

MODEL_MAPPING = {
    "Gemma 3": LLAMACPP_GPU0_URL,
    "Phi 4": LLAMACPP_GPU1_URL,
}
MODEL_OPTIONS = list(MODEL_MAPPING.keys())


# Model-specific temperature settings
MODEL_TEMPERATURES = {
    "Gemma 3": 1.0,
    "Phi 4": 0.5,
}

DEFAULT_MODEL = "Phi 4"
selected_model = DEFAULT_MODEL
DEFAULT_TEMPERATURE = 0.5
MAX_TOKENS = 2048

# Height of the text areas for input and output.
TEXT_AREA_HEIGHT = 400

# Maximum number of characters for the input text.
# This is way below the context window sizes of the models.
# We can increase this. However, we found that users can work and validate better when we nudge to work with shorter texts.
MAX_CHARS_INPUT = 10_000


USER_WARNING = """Mit der KlartextZH-App kannst du Texte sprachlich vereinfachen. Dazu schicken wir deinen Text an einen KI-Server, den das AFI im Kanton betreibt. Du kannst daher auch vertrauliche Daten eingeben.\n\n Bitte beachte: **KI-Sprachmodelle machen Fehler.** Die App liefert lediglich einen Entwurf. **Überprüfe das Ergebnis immer.**\n\nGib uns jederzeit [Feedback](mailto:patrick.arnecke@statistik.ji.zh.ch)."""


# Constants for the formatting of the Word document that can be downloaded.
FONT_WORDDOC = "Arial"
FONT_SIZE_HEADING = 12
FONT_SIZE_PARAGRAPH = 9
FONT_SIZE_FOOTER = 7
DEFAULT_OUTPUT_FILENAME = "Ergebnis.docx"
ANALYSIS_FILENAME = "Analyse.docx"


# Limits for the understandability score to determine if the text is easy, medium or hard to understand.
LIMIT_HARD = -2
LIMIT_MEDIUM = 0

DATETIME_FORMAT = "%Y-%m-%d %H:%M:%S"


# ---------------------------------------------------------------
# Functions


@st.cache_resource
def get_llm_clients():
    """Create connections to the Llama.cpp servers."""
    clients = {}
    for model_name, base_url in MODEL_MAPPING.items():
        if base_url:
            try:
                clients[model_name] = OpenAI(
                    base_url=f"{base_url}/v1",
                    api_key="not-needed",  # API key is not required for Llama.cpp
                )
            except Exception as e:
                st.error(
                    f"Verbindung zum KI-Server für {model_name} fehlgeschlagen: {e}"
                )
                clients[model_name] = None
        else:
            st.warning(f"URL für Modell {model_name} ist nicht konfiguriert.")
            clients[model_name] = None
    return clients


@st.cache_resource
def get_project_info():
    """Get markdown for project information that is shown in the expander section."""
    with open("utils_expander.md") as f:
        return f.read()


@st.dialog("Details zur App", width="large")
def create_project_info(project_info):
    """Create project info."""
    st.markdown(project_info, unsafe_allow_html=True)


def call_llm(
    clients,
    text,
    model_name=DEFAULT_MODEL,
    system_message=DEFAULT_SYSTEM_MESSAGE,
):
    """Call a Llama.cpp server API using OpenAI SDK for text generation."""
    client = clients.get(model_name)
    if not client:
        return False, f"Verbindung zum KI-Server für {model_name} nicht verfügbar."

    temperature = MODEL_TEMPERATURES.get(model_name, DEFAULT_TEMPERATURE)

    # For Llama.cpp's OpenAI compatible endpoint, the model name in the request
    # is often ignored as the server is loaded with a single model.
    # We pass a dummy value like "local-model".
    model_id = "local-model"

    try:
        stream = client.chat.completions.create(
            model=model_id,
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": text.strip()},
            ],
            temperature=temperature,
            max_tokens=MAX_TOKENS,
            stream=True,
        )
        return True, stream

    except Exception as e:
        if "timeout" in str(e).lower():
            return (
                False,
                f"Zeitüberschreitung: Die Anfrage dauerte zu lange und hat nicht funktioniert. Bitte versuche es erneut.",
            )
        return False, str(e)


def create_download_link(text_input, response, selected_model, time_processed):
    """Create a downloadable Word document and download link of the results."""
    document = Document()

    h1 = document.add_heading("Ausgangstext")
    p1 = document.add_paragraph("\n" + text_input)
    h2 = document.add_heading(f"Vereinfachter Text ({selected_model})")
    p2 = document.add_paragraph(response)
    timestamp = datetime.now().strftime(DATETIME_FORMAT)
    footer = document.sections[0].footer
    footer.paragraphs[
        0
    ].text = f"Erstellt am {timestamp} mit der KlartextZH-App des Kantons Zürich.\nModell: {selected_model}\nVerarbeitungszeit: {time_processed:.1f} Sekunden"

    # Set font for all paragraphs.
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = FONT_WORDDOC

    # Set font size for all headings.
    for paragraph in [h1, h2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_HEADING)

    # Set font size for all paragraphs.
    for paragraph in [p1, p2]:
        for run in paragraph.runs:
            run.font.size = Pt(FONT_SIZE_PARAGRAPH)

    # Set font and font size for footer.
    for run in footer.paragraphs[0].runs:
        run.font.name = "Arial"
        run.font.size = Pt(FONT_SIZE_FOOTER)

    section = document.sections[0]
    section.page_width = Inches(8.27)  # Width of A4 paper in inches
    section.page_height = Inches(11.69)  # Height of A4 paper in inches

    io_stream = io.BytesIO()
    document.save(io_stream)

    # # A download button resets the app. So we use a link instead.
    # https://github.com/streamlit/streamlit/issues/4382#issuecomment-1223924851
    # https://discuss.streamlit.io/t/creating-a-pdf-file-generator/7613?u=volodymyr_holomb

    b64 = base64.b64encode(io_stream.getvalue())
    file_name = DEFAULT_OUTPUT_FILENAME
    caption = "Vereinfachten Text herunterladen"

    download_url = f'<a href="data:application/octet-stream;base64,{b64.decode()}" download="{file_name}">{caption}</a>'
    st.markdown(download_url, unsafe_allow_html=True)


def log_event(
    len_text,
    len_response,
    zix_input,
    zix_output,
    simplification_level,
    selected_model,
    time_processed,
    success,
):
    """
    Log a text simplification event with relevant metrics.

    Args:
        len_text : str
            The word count of the original input text.
        len_response : str
            The word count of the simplified output text.
        zix_input : float
            Understandability score for the input text.
        zix_output : float
            Understandability score for the output text.
        simplification_level : str
            The requested level of simplification.
        selected_model : str
            The model name used for simplification.
        time_processed : float
            The processing time in seconds required to generate the simplified text.
        success : bool
            Whether the simplification operation was successful.

    Returns:
        None
            The function logs the information to the configured logging system.

    """
    log_string = f"{datetime.now().strftime(DATETIME_FORMAT)}"
    log_string += f"\t{len_text}"
    log_string += f"\t{len_response}"
    log_string += f"\t{zix_input}"
    log_string += f"\t{zix_output}"
    log_string += f"\t{simplification_level}"
    log_string += f"\t{selected_model}"
    log_string += f"\t{time_processed:.3f}"
    log_string += f"\t{success}"

    logging.info(log_string)

    # Track metrics for Prometheus
    track_metrics(
        len_text,
        len_response,
        zix_input,
        zix_output,
        simplification_level,
        selected_model,
        time_processed,
        success,
    )


# ---------------------------------------------------------------
# Main App

clients = get_llm_clients()
project_info = get_project_info()

with st.sidebar:
    st.markdown(
        "## 🙋‍♀️ KlartextZH - Sprache einfach vereinfachen", unsafe_allow_html=True
    )
    st.markdown(USER_WARNING, unsafe_allow_html=True)
    if st.button("Details zur App", use_container_width=True):
        create_project_info(project_info)
    # selected_model = st.radio(
    #     "Sprachmodell:",
    #     options=MODEL_OPTIONS,
    #     index=MODEL_OPTIONS.index(DEFAULT_MODEL),
    #     help="Wähle das Sprachmodell (LLM), das für die Vereinfachung verwendet werden soll. Phi 4 ist das Standardmodell. Gemma 3 ist ein gutes allgemeines Modell, das auch gute Ergebnisse liefert.",
    # )

button_cols = st.columns([2, 4])

with button_cols[0]:
    # Disable button if client connection for the selected model failed
    simplify_disabled = clients.get(selected_model) is None
    do_simplification = st.button(
        "Vereinfachen",
        use_container_width=True,
        help="Vereinfacht deinen Ausgangstext.",
        type="primary",
        disabled=simplify_disabled,
    )

with button_cols[1]:
    simplification_level = st.radio(
        label="Grad der Vereinfachung",
        options=["Verständliche Sprache", "Einfache Sprache", "Leichte Sprache"],
        index=1,
        help="**«Verständliche Sprache»** überarbeitet den Text vorsichtiger und zielt auf Sprachniveau B2. **«Einfache Sprache»** folgt den Regeln für Einfache Sprache (B1 bis A2) und **«Leichte Sprache»** folgt den Regeln für Leichte Sprache (A2 bis A1).",
        horizontal=True,
    )

st.markdown("---")

system_messages = {
    "Verständliche Sprache": SYSTEM_MESSAGE_VS,
    "Einfache Sprache": SYSTEM_MESSAGE_ES,
    "Leichte Sprache": SYSTEM_MESSAGE_LS,
}

system_message = system_messages[simplification_level]
# if "Gemma" in selected_model:
#     system_message += ADDITION_GEMMA

source_text = st.text_area(
    "Ausgangstext, den du vereinfachen möchtest",
    height=TEXT_AREA_HEIGHT,
    max_chars=MAX_CHARS_INPUT,
    key="key_textinput",
)

source_text = source_text.strip()

if source_text == "":
    st.info("Bitte gib einen Text ein.")
    st.stop()
else:
    score_source = get_zix(source_text)
    # We add 0 to avoid negative zero.
    score_source_rounded = int(np.round(score_source, 0) + 0)
    cefr_source = get_cefr(score_source)

    if score_source < LIMIT_HARD:
        st.markdown(
            f"Dein Ausgangstext ist **:red[schwer verständlich]** und entspricht etwa dem **:red[Sprachniveau {cefr_source}]**."
        )
    elif score_source >= LIMIT_HARD and score_source < LIMIT_MEDIUM:
        st.markdown(
            f"Dein Ausgangstext ist **:orange[durchschnittlich verständlich]** und entspricht etwa dem **:orange[Sprachniveau {cefr_source}]**."
        )
    else:
        st.markdown(
            f"Dein Ausgangstext ist **:green[gut verständlich]** und entspricht etwa dem **:green[Sprachniveau {cefr_source}]**."
        )

word_count = source_text.split()
if len(word_count) <= 5:
    st.error(
        "Dein Text ist sehr kurz. Die App ist dafür gemacht, ganze Sätze bzw. Absätze zu vereinfachen. Das Ergebnis könnte nicht gut sein. Bitte gib einen regulären Text mit ganzen Sätzen ein."
    )

st.markdown("---")

# Start processing
if do_simplification:
    start_time = time.time()

    with st.spinner("Text wird verarbeitet..."):
        success, stream = call_llm(
            clients=clients,
            text=source_text,
            model_name=selected_model,
            system_message=system_message,
        )

    if success is False:
        error_message = str(stream).replace("'", "")[
            :200
        ]  # Limit length and remove quotes for display
        st.error(
            f"Es ist ein Fehler bei der Abfrage aufgetreten: {error_message}. Bitte versuche es erneut."
        )
        time_processed = time.time() - start_time
        error_message_cleaned_for_log = (
            error_message.replace("\t", " ").replace("\n", " ").replace('"', "")
        )
        log_event(
            len(source_text.split()),
            f"Error from model call. {error_message_cleaned_for_log}",
            score_source,
            None,
            simplification_level,
            selected_model,
            time_processed,
            success,
        )

        st.stop()

    # Display results in UI.
    time_processed = time.time() - start_time

    placeholder = st.empty()
    response = ""
    while True:
        try:
            chunk = next(stream)
            if chunk.choices[0].delta.content is not None:
                chunk_text = chunk.choices[0].delta.content
                # Often the models return the German letter ß. Replace it with the Swiss German equivalent ss.
                # Also remove markdown formatting **bold** and # headings.
                chunk_text = chunk_text.replace("ß", "ss")
                chunk_text = re.sub(r"\*\*", " ", chunk_text)
                chunk_text = re.sub(r"^#{1,7}", "", chunk_text, flags=re.MULTILINE)
                if chunk_text == "":
                    continue
                response += chunk_text
                placeholder.text_area(
                    "Dein vereinfachter Text", value=response, height=TEXT_AREA_HEIGHT
                )
        except StopIteration:
            # Finally, remove leading spaces from each line in the response. We cannot remove these during streaming, since it is not included in the chunk with the markdown markup but in the next chunk.
            response_lines = response.split("\n")
            response = "\n".join(
                line[1:] if line.startswith(" ") else line for line in response_lines
            )
            placeholder.text_area(
                "Dein vereinfachter Text", value=response, height=TEXT_AREA_HEIGHT
            )
            break
        except Exception as e:
            st.error(f"Fehler beim Streamen des Textes: {str(e)}")
            break

    if response != "":
        score_target = get_zix(response)
        score_target_rounded = int(np.round(score_target, 0) + 0)
        cefr_target = get_cefr(score_target)

    if score_target < LIMIT_HARD:
        st.markdown(
            f"Dein vereinfachter Text ist **:red[schwer verständlich]** und entspricht etwa dem **:red[Sprachniveau {cefr_target}]**."
        )
    elif score_target >= LIMIT_HARD and score_target < LIMIT_MEDIUM:
        st.markdown(
            f"Dein Ausgangstext ist **:orange[durchschnittlich verständlich]** und entspricht etwa dem **:orange[Sprachniveau {cefr_target}]**."
        )
    else:
        st.markdown(
            f"Dein vereinfachter Text ist **:green[gut verständlich]** und entspricht etwa dem **:green[Sprachniveau {cefr_target}]**."
        )

    create_download_link(
        st.session_state.key_textinput, response, selected_model, time_processed
    )
    st.caption(f"Verarbeitet in {time_processed:.1f} Sekunden.")

    log_event(
        len(source_text.split()),
        len(response.split()),
        np.round(score_source, 2),
        np.round(score_target, 2),
        simplification_level,
        selected_model,
        time_processed,
        success,
    )
    st.stop()
